"""
Document Loader — يدعم PDF, DOCX, Excel, TXT, MD
"""
import re
import os
import tempfile
from pathlib import Path
from typing import Tuple


def clean_text(text: str) -> str:
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[^\S\n]+', ' ', text)
    lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 2]
    return '\n'.join(lines)


def load_pdf(path: str) -> str:
    import PyPDF2
    parts = []
    with open(path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            if t and t.strip():
                parts.append(f"[صفحة {i+1}]\n{t}")
    return '\n\n'.join(parts)


def load_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def load_excel(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[ورقة: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_text = ' | '.join(str(v) for v in row if v is not None)
            if row_text.strip():
                parts.append(row_text)
    return '\n'.join(parts)


def load_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def load_from_bytes(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        if suffix == '.pdf':
            text = load_pdf(tmp_path)
        elif suffix in ('.docx', '.doc'):
            text = load_docx(tmp_path)
        elif suffix in ('.xlsx', '.xls'):
            text = load_excel(tmp_path)
        else:
            text = file_bytes.decode('utf-8', errors='ignore')
        return clean_text(text), filename
    finally:
        os.unlink(tmp_path)
